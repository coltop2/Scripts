import os
import math
from docx import Document
from docx.shared import Inches

# Prompt the user for a directory path
new_dir_name = input("Enter the path where the images are located: ")

# Check if the entered path exists
if not os.path.exists(new_dir_name):
    print("Invalid path. Please enter a valid directory path.")
else:
    # Check if Template.docx exists in the specified directory
    template_file_path = "C:\\Users\\coltop2\\OneDrive - kochind.com\\Documents\\Software Solutions Group\\CMPBOT\\Template.docx"
    if not os.path.exists(template_file_path):
        print("Template.docx not found in the specified directory.")
    else:
        doc = Document(template_file_path)
        images = [image for image in os.listdir(new_dir_name) if image.endswith('.png')]
        table = doc.add_table(rows=math.ceil(len(images) / 2), cols=2)
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                paragraph = cell.add_paragraph()
                words = paragraph.add_run()
                if images:
                    words.add_picture(os.path.join(new_dir_name, images[0]), width=Inches(3.5), height=Inches(2.25))
                    images.pop(0)

        # Prompt the user for the report file name
        report_file_name = input("Enter the name for the report file (without extension): ")
        if not report_file_name:
            report_file_name = "report"  # Default name if the user doesn't provide one

        # Save the document in the specified directory with the user-provided file name
        output_file_path = os.path.join(new_dir_name, f"{report_file_name}.docx")
        doc.save(output_file_path)
        print(f"Report generated and saved at: {output_file_path}")
