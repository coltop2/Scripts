import os
from os import getcwd, listdir, chdir, path
import os.path as osp
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def rename_files(png_folder):
    for entry in os.listdir(png_folder):
        if entry.lower().endswith('.png'):
            if '_SDD21' in entry:
                new_name = entry.replace('_SDD21', f'_01_SDD21')
            elif 'SCD21-SDD21' in entry:
                new_name = entry.replace('SCD21-SDD21', f'02_SCD21-SDD21')
            else:
                new_name = entry
            # Rename the file
            old_path = os.path.join(png_folder, entry)
            new_path = os.path.join(png_folder, new_name)
            os.rename(old_path, new_path)

def main():
    ROOT = os.getcwd()
    template_path = "\\\\ma2vwfp01\\HPC_Group\\Temporary_Share\\Colton Petty\\Comprobot report template\\Template.docx"
    document = Document(template_path)

    png_folder = input("Enter the path to the folder containing PNG files: ")

    # Rename files
    rename_files(png_folder)

    # Add pictures to the document
    i = 0
    for entry in sorted(os.listdir(png_folder)):
        if entry.lower().endswith('.png'):
            document.add_picture(os.path.join(png_folder, entry), height=Inches(3.8))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('')

    output_file_name = input("Enter the desired name for the output docx file (without extension): ")
    output_path = input("Enter the path to save the output docx file: ")
    output_file_path = os.path.join(output_path, output_file_name + '.docx')

    document.save(output_file_path)
    #os.startfile(output_file_path)

if __name__ == '__main__':
    main()
